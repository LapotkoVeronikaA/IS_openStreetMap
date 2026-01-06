# app/organizations/routes.py
import os
import shutil
import uuid
import io
import re
from flask import render_template, request, redirect, url_for, flash, current_app, jsonify, send_file
from sqlalchemy import or_
from werkzeug.utils import secure_filename
from app.models import Organization, GenericDirectoryItem
from app.extensions import db
from app.utils import permission_required_manual, log_user_activity, geocode_location
from . import organizations_bp

from docx import Document
from docx.shared import Inches, Pt
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'pdf'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_files_for_org(org_id, subfolder):
    file_urls = []
    upload_folder = os.path.join(current_app.static_folder, 'uploads', 'organizations', str(org_id), subfolder)
    if os.path.exists(upload_folder):
        for filename in sorted(os.listdir(upload_folder)):
            if '.' in filename:
                file_urls.append(url_for('static', filename=f'uploads/organizations/{org_id}/{subfolder}/{filename}'))
    return file_urls

def sanitize_filename(text):
    return re.sub(r'[\\/*?:"<>|]', "", text).replace(" ", "_")

@organizations_bp.route('/')
@permission_required_manual('view_organizations')
def index():
    # Параметры поиска и фильтрации
    query = request.args.get('q', '').strip()
    org_type = request.args.get('type', '').strip()
    
    # Справочник типов для фильтра
    all_types = db.session.query(Organization.org_type).distinct().all()
    all_types = [t[0] for t in all_types if t[0]]

    is_searching = any([query, org_type])

    if is_searching:
        # Умный поиск по нескольким полям
        search_filter = Organization.query
        if query:
            search_filter = search_filter.filter(or_(
                Organization.name.ilike(f'%{query}%'),
                Organization.legal_name.ilike(f'%{query}%'),
                Organization.location.ilike(f'%{query}%'),
                Organization.head_of_organization.ilike(f'%{query}%'),
                Organization.notes.ilike(f'%{query}%')
            ))
        if org_type:
            search_filter = search_filter.filter(Organization.org_type == org_type)
        
        results = search_filter.order_by(Organization.name).all()
        return render_template('organizations_index.html', results=results, is_searching=True, all_types=all_types)
    
    # Если поиска нет — показываем дерево
    root_orgs = Organization.query.filter_by(parent_id=None).order_by(Organization.name).all()
    return render_template('organizations_index.html', root_orgs=root_orgs, is_searching=False, all_types=all_types)

@organizations_bp.route('/export/search_results')
@permission_required_manual('view_organizations')
def export_search_results():
    """Экспорт результатов текущего поиска в Excel."""
    query = request.args.get('q', '').strip()
    org_type = request.args.get('type', '').strip()

    search_filter = Organization.query
    if query:
        search_filter = search_filter.filter(or_(
            Organization.name.ilike(f'%{query}%'),
            Organization.legal_name.ilike(f'%{query}%'),
            Organization.location.ilike(f'%{query}%'),
            Organization.head_of_organization.ilike(f'%{query}%')
        ))
    if org_type:
        search_filter = search_filter.filter(Organization.org_type == org_type)
    
    results = search_filter.all()

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Результаты поиска"

    headers = ["Название", "Юр. лицо", "Тип", "Адрес", "Руководитель", "Телефон", "Email"]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")

    for org in results:
        ws.append([org.name, org.legal_name, org.org_type, org.location, org.head_of_organization, org.main_phone, org.main_email])

    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return send_file(stream, as_attachment=True, download_name="Registry_Search_Results.xlsx")

@organizations_bp.route('/<int:org_id>')
@permission_required_manual('view_organizations')
def view_org(org_id):
    org = Organization.query.get_or_404(org_id)
    photos = get_files_for_org(org_id, 'photos')
    floor_plans = get_files_for_org(org_id, 'floor_plans')
    return render_template('organization_public_view.html', org=org, photos=photos, floor_plans=floor_plans)

@organizations_bp.route('/add', methods=['GET', 'POST'])
@permission_required_manual('manage_organizations')
def add_org():
    all_orgs = Organization.query.order_by(Organization.name).all()
    organization_types = GenericDirectoryItem.query.filter_by(directory_type='organization_types').all()

    if request.method == 'POST':
        name = request.form.get('name')
        if not name:
            flash('Название обязательно.', 'danger')
            return render_template('organization_form.html', org=None, all_orgs=all_orgs, organization_types=organization_types)
        
        parent_id = request.form.get('parent_id')
        parent_id = int(parent_id) if parent_id and parent_id != 'None' else None

        location_str = request.form.get('location')
        lat, lon = geocode_location(location_str) if location_str else (None, None)

        new_org = Organization(
            name=name,
            legal_name=request.form.get('legal_name'),
            org_type=request.form.get('org_type'),
            parent_id=parent_id,
            location=location_str,
            head_of_organization=request.form.get('head_of_organization'),
            head_position=request.form.get('head_position'),
            website=request.form.get('website'),
            main_phone=request.form.get('main_phone'),
            main_email=request.form.get('main_email'),
            notes=request.form.get('notes'),
            latitude=lat,
            longitude=lon
        )
        
        contacts = []
        for n, p, ph in zip(request.form.getlist('contact_full_name'), request.form.getlist('contact_position'), request.form.getlist('contact_phone')):
            if n.strip(): contacts.append({"full_name": n, "position": p, "phone": ph})
        new_org.set_contacts(contacts)

        db.session.add(new_org)
        db.session.commit()

        for subfolder, file_list in [('photos', request.files.getlist('photos')), ('floor_plans', request.files.getlist('floor_plans'))]:
            if file_list and file_list[0].filename:
                path = os.path.join(current_app.static_folder, 'uploads', 'organizations', str(new_org.id), subfolder)
                os.makedirs(path, exist_ok=True)
                for file in file_list:
                    if file and allowed_file(file.filename):
                        file.save(os.path.join(path, f"{uuid.uuid4()}.{file.filename.rsplit('.', 1)[1].lower()}"))
        
        log_user_activity(f"Добавлена организация: {new_org.name}", "Organization", new_org.id)
        flash('Организация успешно добавлена.', 'success')
        return redirect(url_for('organizations.index'))
            
    return render_template('organization_form.html', org=None, all_orgs=all_orgs, organization_types=organization_types)

@organizations_bp.route('/edit/<int:org_id>', methods=['GET', 'POST'])
@permission_required_manual('manage_organizations')
def edit_org(org_id):
    org = Organization.query.get_or_404(org_id)
    all_orgs = Organization.query.filter(Organization.id != org_id).order_by(Organization.name).all()
    organization_types = GenericDirectoryItem.query.filter_by(directory_type='organization_types').all()

    if request.method == 'POST':
        org.name = request.form.get('name')
        parent_id = request.form.get('parent_id')
        org.parent_id = int(parent_id) if parent_id and parent_id != 'None' else None

        new_loc = request.form.get('location')
        if new_loc != org.location:
             org.latitude, org.longitude = geocode_location(new_loc)
        org.location = new_loc

        org.legal_name = request.form.get('legal_name')
        org.org_type = request.form.get('org_type')
        org.head_of_organization = request.form.get('head_of_organization')
        org.head_position = request.form.get('head_position')
        org.website = request.form.get('website')
        org.main_phone = request.form.get('main_phone')
        org.main_email = request.form.get('main_email')
        org.notes = request.form.get('notes')
        
        contacts = []
        for n, p, ph in zip(request.form.getlist('contact_full_name'), request.form.getlist('contact_position'), request.form.getlist('contact_phone')):
            if n.strip(): contacts.append({"full_name": n, "position": p, "phone": ph})
        org.set_contacts(contacts)

        for subfolder, file_list in [('photos', request.files.getlist('photos')), ('floor_plans', request.files.getlist('floor_plans'))]:
            if file_list and file_list[0].filename:
                path = os.path.join(current_app.static_folder, 'uploads', 'organizations', str(org_id), subfolder)
                os.makedirs(path, exist_ok=True)
                for file in file_list:
                    if file and allowed_file(file.filename):
                        file.save(os.path.join(path, f"{uuid.uuid4()}.{file.filename.rsplit('.', 1)[1].lower()}"))

        db.session.commit()
        log_user_activity(f"Обновлена организация: {org.name}", "Organization", org.id)
        flash('Данные обновлены.', 'success')
        return redirect(url_for('organizations.index'))
            
    photos = get_files_for_org(org_id, 'photos')
    floor_plans = get_files_for_org(org_id, 'floor_plans')
    return render_template('organization_form.html', org=org, all_orgs=all_orgs, photos=photos, floor_plans=floor_plans, organization_types=organization_types)

@organizations_bp.route('/delete/<int:org_id>', methods=['POST'])
@permission_required_manual('manage_organizations')
def delete_org(org_id):
    org = Organization.query.get_or_404(org_id)
    upload_folder = os.path.join(current_app.static_folder, 'uploads', 'organizations', str(org_id))
    if os.path.exists(upload_folder): shutil.rmtree(upload_folder)
    db.session.delete(org)
    db.session.commit()
    flash('Организация удалена.', 'success')
    return redirect(url_for('organizations.index'))

@organizations_bp.route('/delete_file/<int:org_id>/<string:subfolder>/<string:filename>', methods=['POST'])
@permission_required_manual('manage_organizations')
def delete_file(org_id, subfolder, filename):
    filename = secure_filename(filename)
    file_path = os.path.join(current_app.static_folder, 'uploads', 'organizations', str(org_id), subfolder, filename)
    if os.path.exists(file_path):
        os.remove(file_path)
        return jsonify({'success': True})
    return jsonify({'success': False}), 404

@organizations_bp.route('/export/docx/<int:org_id>')
def export_docx(org_id):
    org = Organization.query.get_or_404(org_id)
    doc = Document()
    header = doc.add_heading(f'Информационная карточка объекта', 0)
    header.alignment = 1
    doc.add_heading('1. Общие сведения', level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    data = [('Наименование:', org.name), 
            ('Юридическое лицо:', org.legal_name or '-'), 
            ('Тип объекта:', org.org_type or '-'), 
            ('Адрес местонахождения:', org.location or '-'), 
            ('Руководитель:', f"{org.head_position or 'Руководитель'}: {org.head_of_organization or '-'}")]
    for key, value in data:
        row = table.add_row().cells
        row[0].text = key
        row[0].paragraphs[0].runs[0].bold = True
        row[1].text = str(value)
    doc.add_heading('2. Контактные лица и штат', level=1)
    contacts = org.get_contacts()
    if contacts:
        c_table = doc.add_table(rows=1, cols=3)
        c_table.style = 'Table Grid'
        hdrs = c_table.rows[0].cells
        hdrs[0].text, hdrs[1].text, hdrs[2].text = 'ФИО', 'Должность', 'Телефон'
        for c in contacts:
            row = c_table.add_row().cells
            row[0].text, row[1].text, row[2].text = c['full_name'], c['position'], c['phone']
    doc.add_heading('3. Внутренняя структура', level=1)
    children = org.children.all()
    if children:
        for child in children:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{child.name}").bold = True
            p.add_run(f" ({child.org_type or 'подразделение'})")
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return send_file(file_stream, 
                     as_attachment=True, 
                     download_name=f"Report_{sanitize_filename(org.name)}.docx", 
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@organizations_bp.route('/export/xlsx/<int:org_id>')
def export_xlsx(org_id):
    org = Organization.query.get_or_404(org_id)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Общая информация"
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = openpyxl.styles.PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
    ws['A1'], ws['B1'] = "Параметр", "Значение"
    for cell in ['A1', 'B1']: ws[cell].font = header_font; ws[cell].fill = header_fill
    rows = [["ID", org.id], 
            ["Название", org.name], 
            ["Юр. лицо", org.legal_name], 
            ["Тип", org.org_type], 
            ["Адрес", org.location], 
            ["Руководитель", org.head_of_organization], 
            ["Должность рук.", org.head_position], 
            ["Сайт", org.website], 
            ["Email", org.main_email], 
            ["Телефон", org.main_phone]]
    for r in rows: ws.append(r)
    file_stream = io.BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)
    return send_file(file_stream, 
                     as_attachment=True, 
                     download_name=f"Report_{sanitize_filename(org.name)}.xlsx", 
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')